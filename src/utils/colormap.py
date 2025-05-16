from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.shared import Mm
import os

# Utility conversions
dpi = 96
mm_to_px = lambda mm: int(mm * dpi / 25.4)

# Candidate Japanese fonts (common on many Linux distros)
font_candidates = [
    "/usr/share/fonts/truetype/noto/NotoSansCJK-Regular.ttc",
    "/usr/share/fonts/truetype/noto/NotoSansCJKJP-Regular.otf",
    "/usr/share/fonts/truetype/takao-gothic/TakaoPGothic.ttf",
    "/usr/share/fonts/truetype/ipafont-gothic/ipagp.ttf",
    "/usr/share/fonts/opentype/ipaexfont-gothic/ipaexg.ttf",
]

font_path = None
for cand in font_candidates:
    if os.path.exists(cand):
        font_path = cand
        break

if font_path is None:
    # Fallback to DejaVuSans (may still fail for JP glyphs)
    font_path = "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf"

# Create pattern function
def create_patterned_rect(size, line_color):
    w, h = size
    img = Image.new("RGBA", (w, h), (0, 0, 0, 0))
    d = ImageDraw.Draw(img)
    spacing = 8
    for x in range(-h, w, spacing):
        d.line([(x, 0), (x + h, h)], fill=line_color + (60,), width=4)
    return img

# Build diagram image
canvas_w, canvas_h = 800, 400
canvas = Image.new("RGBA", (canvas_w, canvas_h), "white")

font_small = ImageFont.truetype(font_path, 24)

# Rectangle
rect_w, rect_h = mm_to_px(25), mm_to_px(30)
rect_img = Image.new("RGBA", (rect_w, rect_h), (255, 0, 0, 40))
rect_img = Image.alpha_composite(rect_img, create_patterned_rect((rect_w, rect_h), (255, 0, 0)))
draw = ImageDraw.Draw(rect_img)
draw.rectangle([(0, 0), (rect_w - 1, rect_h - 1)], outline=(200, 0, 0), width=4)
bbox = draw.textbbox((0, 0), "しかく", font=font_small)
w, h = bbox[2] - bbox[0], bbox[3] - bbox[1]
draw.text(((rect_w - w) / 2, (rect_h - h) / 2), "しかく", font=font_small, fill="black")
rect_img = rect_img.rotate(25, expand=True)
canvas.alpha_composite(rect_img, (50, 100))

# Triangle
tri_side = mm_to_px(45)
tri_img = Image.new("RGBA", (tri_side, tri_side), (0, 0, 0, 0))
tri_img = Image.alpha_composite(tri_img, create_patterned_rect((tri_side, tri_side), (255, 215, 0)))
t_draw = ImageDraw.Draw(tri_img)
pts = [(tri_side / 2, 0), (0, tri_side - 1), (tri_side - 1, tri_side - 1)]
t_draw.polygon(pts, outline=(200, 180, 0), width=4)
bbox = t_draw.textbbox((0, 0), "さんかく", font=font_small)
w, h = bbox[2] - bbox[0], bbox[3] - bbox[1]
t_draw.text((tri_side / 2 - w / 2, tri_side * 0.55), "さんかく", font=font_small, fill="black")
canvas.alpha_composite(tri_img, (300, 80))

# Circle
circle_d = mm_to_px(35)
circ_img = Image.new("RGBA", (circle_d, circle_d), (0, 0, 0, 0))
circ_img = Image.alpha_composite(circ_img, create_patterned_rect((circle_d, circle_d), (0, 0, 255)))
c_draw = ImageDraw.Draw(circ_img)
c_draw.ellipse([(0, 0), (circle_d - 1, circle_d - 1)], outline=(0, 0, 180), width=4)
bbox = c_draw.textbbox((0, 0), "えん", font=font_small)
w, h = bbox[2] - bbox[0], bbox[3] - bbox[1]
c_draw.text((circle_d / 2 - w / 2, circle_d / 2 - h / 2), "えん", font=font_small, fill="black")
canvas.alpha_composite(circ_img, (600, 110))

img_path = "/mnt/data/shapes_jp.png"
canvas.convert("RGB").save(img_path)

# Build Word document
doc = Document()
doc.add_heading("課題1：図形描画サンプル（日本語文字修正版）", level=1)
doc.add_paragraph("以下は Word 内に配置した図形のイメージ例です（自動生成）。")
doc.add_picture(img_path, width=Mm(160))

doc.add_page_break()
doc.add_heading("課題2：自己紹介", level=1)
doc.add_paragraph("氏名：上村")
doc.add_paragraph("所属：明治大学 電気電子工学科 4年")
doc.add_paragraph("研究分野：AI・機械学習・コンピュータビジョン")
doc.add_paragraph("趣味：ソフトテニス、プログラミング、読書")
doc.add_paragraph("（写真をここに挿入）")

doc_path = "/mnt/data/Assignment_Completed_JP.docx"
doc.save(doc_path)

print(f"Japanese-support Word file created at: {doc_path}")

